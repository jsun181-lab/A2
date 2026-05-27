from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "report"
OUTPUT = REPORT / "Report_Junsi Sun.docx"
SYSTEM_DESIGN = REPORT / "system-design.png"

REPO_LINK = "https://github.com/jsun181-lab/A2.git"

INK = RGBColor(23, 33, 43)
MUTED = RGBColor(82, 98, 113)
BLUE = RGBColor(46, 116, 181)
TEAL = RGBColor(20, 118, 109)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.58)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(9.15)
    normal.paragraph_format.space_after = Pt(3.1)
    normal.paragraph_format.line_spacing = 1.03

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(8.9)
        style.paragraph_format.space_after = Pt(2.0)
        style.paragraph_format.line_spacing = 1.02

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("COMPSCI 767 Assignment 2 | DataInsight Agent")
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor(120, 132, 146)


def set_cell_margins(cell, top=55, start=80, bottom=55, end=80):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_table_borders(table, color="DADCE0"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(title)
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = INK

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3.5)
    run = p.add_run(subtitle)
    run.font.name = "Calibri"
    run.font.size = Pt(9.2)
    run.font.color.rgb = MUTED


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4.0 if level == 1 else 2.5)
    p.paragraph_format.space_after = Pt(1.8)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.bold = True
    run.font.size = Pt(12.5 if level == 1 else 10.0)
    run.font.color.rgb = BLUE if level == 1 else TEAL


def add_label_value(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2.8)
    label_run = p.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.color.rgb = INK
    value_run = p.add_run(value)
    value_run.font.color.rgb = RGBColor(52, 76, 154)


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3.1)
    p.paragraph_format.line_spacing = 1.03
    run = p.add_run(text)
    run.font.size = Pt(9.15)
    run.font.color.rgb = INK


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.first_line_indent = Inches(-0.08)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.24)
        p.paragraph_format.first_line_indent = Inches(-0.08)
        p.add_run(item)


def add_text_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = True
    set_table_borders(table)
    for row_index, (label, detail) in enumerate(rows):
        left = table.cell(row_index, 0)
        right = table.cell(row_index, 1)
        set_cell_margins(left)
        set_cell_margins(right)
        if row_index == 0:
            set_cell_shading(left, "F2F4F7")
            set_cell_shading(right, "F2F4F7")
        label_run = left.paragraphs[0].add_run(label)
        label_run.bold = True
        label_run.font.size = Pt(8.6)
        label_run.font.color.rgb = TEAL
        detail_run = right.paragraphs[0].add_run(detail)
        detail_run.font.size = Pt(8.55)
        detail_run.font.color.rgb = INK


def build():
    REPORT.mkdir(parents=True, exist_ok=True)
    if not SYSTEM_DESIGN.exists():
        raise FileNotFoundError(f"Missing system design diagram: {SYSTEM_DESIGN}")

    doc = Document()
    style_document(doc)

    add_title(doc, "DataInsight Agent", "COMPSCI 767 Assignment 2 - Intelligent CSV Cleaning and Analysis Agent Prototype")
    add_label_value(doc, "GitHub repository", REPO_LINK)
    add_heading(doc, "System Design Diagram", 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2.5)
    p.add_run().add_picture(str(SYSTEM_DESIGN), width=Inches(6.55))

    add_heading(doc, "Design Explanation", 1)
    add_body(
        doc,
        "DataInsight Agent is a browser-based intelligent software agent that helps a user turn a messy CSV file into a cleaner dataset, useful charts, and short analysis findings. The system is intentionally small, local, and reproducible, but it still follows the core agent pattern from the course: it perceives input, reasons over the current state, selects actions toward a goal, executes those actions, and remembers the result of the run.",
    )
    add_body(
        doc,
        "The agent's goal is to improve a tabular dataset enough for quick exploratory analysis. The user can paste CSV text, upload a CSV file, or load the sample sales dataset. The selected cleaning mode acts as the user's goal constraint: conservative mode avoids changing numeric meaning, balanced mode applies practical low-risk repairs, and aggressive mode allows stronger transformations such as numeric imputation and outlier capping.",
    )
    add_text_table(
        doc,
        [
            ("Agent stage", "Role in the prototype"),
            ("Perception", "The CSV parser reads headers and rows, infers data types, counts missing values, detects duplicate rows, identifies category inconsistencies, finds numeric outliers, and flags possible sensitive columns."),
            ("Decision", "The planner converts the perceived profile into a cleaning plan and an analysis plan. It chooses actions based on dataset issues, available column types, the selected goal, and the user's risk tolerance."),
            ("Action", "The executor removes duplicates, fills or flags missing values, standardizes categories, caps or labels outliers when appropriate, generates charts, summarizes findings, and exports a cleaned CSV."),
            ("Memory and safety", "The browser stores recent run checkpoints in localStorage. The interface warns users to use non-sensitive data only, and the decision output marks risk when cleaning may affect interpretation."),
        ],
    )
    add_heading(doc, "Why This Is an Intelligent Agent", 1)
    add_bullets(
        doc,
        [
            "It is not only a static dashboard: the output changes according to the CSV content, selected cleaning mode, and detected analysis opportunities.",
            "It uses perception to build an internal representation of the environment, including data quality score, missingness, duplicates, outliers, and column semantics.",
            "It makes decisions by selecting actions that move the dataset toward the user goal of clean, explainable exploratory analysis.",
            "It takes visible actions by transforming data, producing charts, enabling CSV export, and writing a memory checkpoint for future reference.",
            "It includes safety boundaries by keeping all processing local and by making risky transformations visible to the user.",
        ],
    )

    doc.add_page_break()
    add_title(doc, "How the System Works", "Detailed walkthrough of the running prototype")
    add_heading(doc, "Runtime Workflow", 1)
    add_numbered(
        doc,
        [
            "The user opens the local web app and either loads the sample CSV or provides their own CSV text. This is the agent's environment input.",
            "When Run Agent is clicked, the perception module parses the CSV and creates a profile. In the sample dataset it identifies 24 rows, 6 columns, a duplicate row, one missing Region value, inconsistent category casing, and revenue outliers.",
            "The decision module compares those issues with the selected cleaning mode. In balanced mode it chooses practical repairs such as removing duplicates, filling unknown categories, standardizing category names, and flagging outliers instead of silently changing extreme numeric values.",
            "The action module applies the plan to the rows and recalculates the quality score. In the sample run, the quality score improves from 94/100 to 98/100, which demonstrates progress toward the cleaning goal.",
            "The analysis module inspects available column types and creates charts such as bar, donut, histogram, line, and scatter views. These charts are not hard-coded screenshots; they are generated from the cleaned data produced by the agent.",
            "Finally, the app stores a memory checkpoint containing the file name, mode, before/after quality score, and headline finding. This gives the prototype state across runs without requiring an external database or API key.",
        ],
    )
    add_heading(doc, "Expected Screenshots and Evidence", 1)
    add_body(
        doc,
        "The first screenshot for the report should show the main result screen after clicking Load Sample and Run Agent. The important evidence on that screen is the perceive-decide-act pipeline: metrics for rows, columns, before quality, and after quality; an Agent Decision card showing balanced cleaning and low risk; a Data Profile panel listing detected column types and missing values; and a Cleaning Strategy panel listing the selected actions.",
    )
    add_body(
        doc,
        "The second screenshot should show the generated chart area and memory section. This view demonstrates that the agent acted on the cleaned data rather than merely describing it. The chart panel shows multiple chart types chosen from the data profile, while the memory panel shows the run saved as a checkpoint. Together these screenshots provide evidence for perception, decision making, action, and memory.",
    )
    add_heading(doc, "Implementation Notes", 1)
    add_bullets(
        doc,
        [
            "The project is implemented with plain HTML, CSS, and JavaScript so it can run locally through python -m http.server 8765.",
            "Core logic is separated from the UI: CSV parsing is handled in csv.js, agent planning and execution are handled in agent.js, and browser rendering is handled in app.js.",
            "No external LLM or cloud API is required. This design makes the prototype easy to reproduce and avoids accidental data sharing.",
            "The test suite checks CSV parsing, quality improvement, aggressive outlier handling, and memory checkpoint behavior.",
        ],
    )
    add_heading(doc, "Limitations and Future Improvements", 1)
    add_body(
        doc,
        "The current prototype is rule-based, so it is transparent and reliable for the sample task but less flexible than an LLM-powered agent. Future versions could add natural-language goals, richer anomaly explanations, undo/redo for cleaning actions, and an optional LLM summary step. The main safety improvement would be a stronger sensitive-data detector before any advanced analysis is allowed.",
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
